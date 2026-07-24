"""
Assemble the full workshop write-up as a Word document, using every figure in
figures/ plus tables of the real pipeline numbers. This is a *report builder*,
not a modeling script: it imports nothing from fedconformal except to read
figure files and does not recompute results (those live in the other scripts
and notebooks; the numbers quoted in the prose below were read off their
output and are refreshed by hand if the pipeline numbers change).

Structure: Introduction -> Challenges -> Motivation -> The
Difficulty of Working With Real Clinical Data -> Our Dataset -> Methods of
Interoperability (standards survey) -> Why Conformal Prediction
Helps -> The Mathematics -> Pipeline Walkthrough -> Appendix.

Run:  python scripts/generate_report.py
Writes: report/Beyond_Interoperability_Report.docx
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), "..")
FIG = os.path.join(ROOT, "figures")
OUT_DIR = os.path.join(ROOT, "report")
OUT = os.path.join(OUT_DIR, "Beyond_Interoperability_Report.docx")
os.makedirs(OUT_DIR, exist_ok=True)

INK = RGBColor(0x1a, 0x1a, 0x1a)
ACCENT = RGBColor(0x2a, 0x78, 0xd6)
MUTED = RGBColor(0x60, 0x60, 0x60)


# ----------------------------------------------------------------------------
# Small helpers
# ----------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(size)


def add_numbered(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(it)
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(6)


def add_image(doc, path, caption, width=6.2):
    full = os.path.join(FIG, path)
    if not os.path.exists(full):
        add_para(doc, f"[missing figure: {path}]", italic=True, color=MUTED)
        return
    doc.add_picture(full, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED
    cap.paragraph_format.space_after = Pt(16)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, widths=None, header_fill="2A78D6"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], header_fill)
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_table_caption(doc, text):
    add_para(doc, text, bold=True, size=10.5)


def add_math(doc, parts, size=13.5):
    """A formula rendered with real Unicode math symbols (alpha, hats via a
    combining circumflex, ceiling/floor brackets, sigma, minus/leq signs) and
    true sub/superscript run formatting in Cambria Math -- as close to a live
    equation as plain docx runs get without hand-building an OMML object.

    ``parts`` is a list of (text, style) pairs, style in
    {"normal", "sub", "sup"}, concatenated left to right on one line.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, kind in parts:
        run = p.add_run(text)
        run.font.name = "Cambria Math"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1c, 0x5c, 0xab)
        if kind == "sub":
            run.font.subscript = True
        elif kind == "sup":
            run.font.superscript = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F0F4FA")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    return p


def page_break(doc):
    doc.add_page_break()


# ----------------------------------------------------------------------------
# Build
# ----------------------------------------------------------------------------

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = INK

# ---- Title page -------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Beyond Interoperability: Hands-On Federated ML for "
                    "Research Data Curation Infrastructure")
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Technical Companion: Conformal Prediction for Site-Level Heterogeneity\n"
                  "in Federated Clinical Data")
run.font.size = Pt(15)
run.font.color.rgb = INK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("A technical report and workshop companion for the fedconformal-clinical toolkit\n"
                   "UCI Heart Disease federation · 4 hospitals · 920 patients")
run.font.size = Pt(11.5)
run.italic = True
run.font.color.rgb = MUTED

# Author / affiliation, as given by the workshop proposal. Venue/date was not
# provided, so that field alone stays a flagged placeholder rather than invented.
PLACEHOLDER = RGBColor(0xc0, 0x3a, 0x2b)
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_before = Pt(28)
run = authors.add_run("Arnab Mukherjee")
run.font.size = Pt(12.5)
run.bold = True
run.font.color.rgb = INK
run = authors.add_run("\nOklahoma State University — Center for Health Science, Tulsa")
run.font.size = Pt(11)
run.font.color.rgb = INK

venue = doc.add_paragraph()
venue.alignment = WD_ALIGN_PARAGRAPH.CENTER
venue.paragraph_format.space_before = Pt(10)
run = venue.add_run("[Workshop / venue name and date — add here]")
run.font.size = Pt(11)
run.font.color.rgb = PLACEHOLDER

page_break(doc)

# ---- Abstract -------------------------------------------------------------
# Verbatim from the workshop proposal -- not paraphrased.
add_heading(doc, "Abstract", level=1)
add_para(doc,
    "Interoperability—the ability of systems to securely exchange, interpret, and use data "
    "without human intervention—has become a foundation of modern research data curation "
    "infrastructure. Yet interoperability alone does not establish whether data from different sites "
    "encode equivalent underlying phenomena. When clinical data exhibits a distributional shift "
    "across institutions, models trained at one site may fail silently at another. Confident but "
    "wrong is the worst possible failure mode, particularly when models are deployed across "
    "organizations. This is the foundational curation challenge this workshop addresses. Research "
    "software without borders requires not just that data can cross institutional lines, but that "
    "what the data means can be trusted on the other side. Site-level measurement heterogeneity "
    "determines whether data generated across diverse sources can be meaningfully compared, "
    "interpreted, and trusted as part of a shared research infrastructure. We complement "
    "interoperability-focused approaches by developing methods to identify and quantify site-level "
    "heterogeneity before data and models are integrated, thereby distinguishing true underlying "
    "variation from measurement-induced differences and establishing a more reliable foundation for "
    "curation in cross-site federated research.",
    size=11)
add_para(doc,
    "This workshop translates that framework into practice. Participants will engage in hands-on "
    "Python exercises using conformal prediction under simulated site-level heterogeneity, working "
    "directly with federated ML tools designed for research data curation pipelines. Attendees will "
    "leave with reusable code, a conceptual framework for distinguishing distributional shift from "
    "measurement-induced heterogeneity, and practical uncertainty quantification tools they can "
    "deploy within their own cross-institutional curation infrastructure.",
    size=11)

page_break(doc)

# ---- Table of contents ---------------------------------------------------
add_heading(doc, "Contents", level=1)
toc_items = [
    "1. Introduction",
    "2. Challenges",
    "3. Motivation",
    "4. The Difficulty of Working With Real Clinical Data",
    "5. Our Dataset: The UCI Heart Disease Federation",
    "6. Methods of Interoperability: Choosing a Standard",
    "7. Why Conformal Prediction Helps With Site-Level Heterogeneity",
    "8. The Mathematics, and What Every Piece of Code Produces",
    "9. Pipeline Walkthrough and Justification: Is This Enough?",
    "References",
    "Appendix: Full Figure Index",
]
for it in toc_items:
    add_para(doc, it, size=12)
page_break(doc)

# ==========================================================================
# SECTION 1 -- INTRODUCTION
# ==========================================================================
add_heading(doc, "1. Introduction", level=1)
add_para(doc,
    "Interoperability means two systems can exchange and read each other's data without a human "
    "translating in between. It says nothing about whether the data means the same thing once it "
    "arrives. A hospital's cholesterol field and another hospital's cholesterol field can sit in the "
    "identical column, in the identical format, and still describe two different measurement "
    "processes -- one hospital records it, another does not. That gap between \"the data moved "
    "successfully\" and \"the data can be trusted the same way everywhere\" is the problem this "
    "project addresses.")
add_para(doc,
    "We use conformal prediction -- a statistical wrapper that turns any machine-learning model's "
    "output into an honest, checkable range of possible answers -- as a diagnostic instrument for "
    "that gap. A conformal predictor comes with a guarantee: at the site where it was calibrated, "
    "its prediction ranges will contain the true answer at a chosen rate (say, 90% of the time), no "
    "matter what model sits underneath it. That guarantee is a promise about *one* site. The moment "
    "you deploy the same predictor at a second site, the promise is only as good as the assumption "
    "that the second site's data looks statistically like the first. When it does not -- when there "
    "is real site-level heterogeneity -- the guarantee breaks, visibly and measurably, and conformal "
    "prediction is precisely the tool that lets you see it break instead of finding out the hard way "
    "in production.")
add_para(doc,
    "This report walks through the full toolkit built to make that idea concrete and hands-on: the "
    "UCI Heart Disease dataset loaded from its four original hospital sites (Cleveland, Hungary, "
    "Switzerland, the V.A. Long Beach), a from-scratch federated-learning simulator, a from-scratch "
    "conformal-prediction library, and a battery of diagnostics that separate genuine population "
    "differences (\"this hospital really does see sicker patients\") from measurement artefacts "
    "(\"this hospital simply never recorded this lab value\"). Every figure the pipeline produces -- "
    "33 in total -- is reproduced here with a plain-language explanation of what it shows and why it "
    "matters, alongside the mathematics behind the method and an honest assessment of what is solid "
    "and what could still be added.")
add_para(doc,
    "The rest of this report is organized to build the argument from the ground up rather than jump "
    "straight to results. Section 2 lays out the concrete challenges a cross-institution curation "
    "pipeline has to solve. Section 3 explains why this is a live, fast-moving problem right now, not "
    "a settled one. Section 4 confronts a practical reality that most technical write-ups skip over "
    "entirely: why real multi-institution clinical data is so hard to obtain in the first place, and "
    "what that implies for anyone trying to build or teach this kind of pipeline. Section 5 introduces "
    "the specific dataset this toolkit uses and why it is an unusually good fit. Section 6 asks a "
    "question that is easy to leave implicit: given that interoperability itself is a solved technical "
    "problem with multiple competing standards, which standard actually fits this use case, and what "
    "does adopting one buy you (and not buy you)? Sections 7 through 9 then get into the technical "
    "core -- the method, the mathematics, and a full walkthrough of the pipeline -- and the report "
    "closes with an honest verdict on what is solid today and what a production deployment should add "
    "next.")

page_break(doc)

# ==========================================================================
# SECTION 2 -- CHALLENGES
# ==========================================================================
add_heading(doc, "2. Challenges", level=1)
add_para(doc,
    "Before describing the toolkit, it is worth naming the problem precisely. Five distinct "
    "challenges compound to make cross-institution clinical modeling harder than it looks from the "
    "outside, and conflating them is one of the easiest ways a curation project goes wrong.")

add_bullets(doc, [
    "Separating true variation from measurement artefact, with no ground truth. Given a feature "
    "that differs across two hospitals, there is no automatic way to know whether the underlying "
    "patients really differ or whether one hospital simply measures that variable differently (or "
    "not at all). Both look identical in a naive summary statistic; only a deliberate, "
    "purpose-built diagnostic tells them apart.",
    "Silent failure. A model's own output gives no warning when it has left the population it was "
    "trained on. A softmax score of 92% looks the same whether it is well-calibrated or not -- the "
    "model does not know, and cannot say, that it is now operating somewhere it has never been "
    "validated.",
    "Data cannot simply be pooled to fix the first two problems. The obvious fix for heterogeneity "
    "-- combine everyone's data, retrain, re-validate -- is usually not legally or practically "
    "available in healthcare (Section 4 explains why in detail). Any solution has to work under the "
    "constraint that patient-level records stay inside their originating institution.",
    "A shared schema does not imply shared meaning. Two institutions can adopt the identical data "
    "standard, the identical column names, the identical value codes, and still encode different "
    "underlying clinical practices -- which tests get ordered, which values get routinely recorded, "
    "which patients get referred there in the first place. Standardization (Section 6) narrows this "
    "gap; it does not close it.",
    "Standard evaluation metrics do not expose any of this. Accuracy, AUC, and even a well-calibrated "
    "pooled coverage number are all computed as one number across the whole validation set. Every "
    "one of them can look healthy while one site, or one patient subgroup, is being served far below "
    "the reported standard. Finding that requires metrics that are explicitly stratified -- by site, "
    "by class, by subgroup -- not just a bigger validation set.",
])
add_para(doc,
    "These five challenges are the reason this report is organized the way it is: Section 6 addresses "
    "the third and fourth (what standardization can and cannot do), and Sections 7 through 9 build and "
    "demonstrate the diagnostic tooling that addresses the first, second and fifth directly.")

page_break(doc)

# ==========================================================================
# SECTION 3 -- MOTIVATION
# ==========================================================================
add_heading(doc, "3. Motivation", level=1)
add_para(doc,
    "None of the challenges in Section 2 are new in principle. What has changed is that three "
    "independent trends are accelerating at the same time, and the gap between them is exactly where "
    "this toolkit sits.")

add_heading(doc, "3.1 Clinical AI is being deployed faster than it is being validated across sites", level=2)
add_para(doc,
    "It is well documented in the clinical-ML literature that models validated at one hospital can "
    "underperform when deployed at another -- differences in patient case-mix, equipment, and "
    "measurement practice are enough to degrade performance even when the input schema is identical. "
    "Regulators have taken notice of the same pattern: the FDA's approach to AI/ML-enabled medical "
    "devices now explicitly includes real-world performance monitoring after deployment -- tracking "
    "for performance drift and monitoring across subpopulations, not stopping at a single pre-market "
    "validation number [FDA-AI-SaMD]. That is precisely the per-site, per-subgroup accountability "
    "this toolkit's diagnostics are built to provide.")

add_heading(doc, "3.2 Federated learning has moved from research curiosity to practical necessity", level=2)
add_para(doc,
    "Federated learning -- training a shared model by exchanging weights instead of data -- was "
    "introduced by McMahan et al. (2017) [McMahan17] and popularized soon after for keyboard-"
    "prediction models on phones (Hard et al., 2018) [Hard18], where centralizing personal keystroke "
    "data was undesirable but not legally prohibitive. In healthcare the same architecture is close to "
    "a practical necessity rather than a stylistic choice, because (as Section 4 details) centralizing "
    "patient-level records across institutions is frequently not something a legal and regulatory "
    "process will approve on any reasonable timeline. Purpose-built federated healthcare benchmarks "
    "such as FLamby (Terrail et al., 2022) [FLamby22] -- which includes the same four-hospital UCI "
    "Heart Disease population used in this report, as its own filtered \"Fed-Heart-Disease\" task "
    "(one of seven cross-silo benchmarks) -- reflect the same constraint this report's simulator is "
    "built around: train without moving the data.")

add_heading(doc, "3.3 Interoperability mandates target structural, not semantic, interoperability", level=2)
add_para(doc,
    "In the United States, the ONC's 21st Century Cures Act Final Rule required certified health IT "
    "developers to expose FHIR-based APIs to their customers [ONC-Cures], and the Trusted Exchange "
    "Framework and Common Agreement (TEFCA) is now an operating nationwide network for health "
    "information exchange, with tens of thousands of participating sites [TEFCA]. In the European "
    "Union, the European Health Data Space Regulation (EU) 2025/327 was adopted in early 2025 and "
    "will progressively apply from 2027 [EHDS25]. These are genuine, significant achievements -- but "
    "each one mandates *structural* interoperability (can the bytes move, in an agreed format) far "
    "more directly than *semantic* interoperability (does the moved value mean the same clinical fact "
    "everywhere; see the HIMSS levels in Section 6.1). As more institutions become technically "
    "interoperable by mandate, the silent semantic gap this report measures becomes more consequential, "
    "not less: it is the remaining barrier standing between \"the data arrived\" and \"the data can be "
    "trusted,\" which is exactly the subject of Section 6.")

add_heading(doc, "3.4 Auditable uncertainty quantification is a natural complement to these mandates", level=2)
add_para(doc,
    "Conformal prediction is one of very few uncertainty-quantification methods that offers a "
    "distribution-free, finite-sample coverage guarantee -- a mathematical property (proved in Section "
    "8) that does not depend on the underlying model being well-specified, Bayesian, or of any "
    "particular architecture. That property is exactly what a per-site, auditable accountability check "
    "needs: a number a reviewer can verify empirically, not one that has to be taken on faith in the "
    "model.")
add_para(doc,
    "Put together: clinical AI deployment, structural interoperability mandates, and the case for "
    "auditable uncertainty each stand on independently verifiable ground, and the space between the "
    "second and third -- interoperable, but not yet verifiably trustworthy -- is the open problem this "
    "report and its accompanying toolkit are built to make tractable.")

add_para(doc,
    "Sources for this section: "
    "[FDA-AI-SaMD] U.S. FDA, \"Artificial Intelligence in Software as a Medical Device\" and the "
    "AI/ML-Based SaMD Action Plan, fda.gov/medical-devices/software-medical-device-samd. "
    "[McMahan17] H. B. McMahan et al., \"Communication-Efficient Learning of Deep Networks from "
    "Decentralized Data,\" AISTATS 2017. "
    "[Hard18] A. Hard et al., \"Federated Learning for Mobile Keyboard Prediction,\" 2018 "
    "(arXiv:1811.03604). "
    "[FLamby22] J. O. du Terrail et al., \"FLamby: Datasets and Benchmarks for Cross-Silo Federated "
    "Learning in Realistic Healthcare Settings,\" 2022 (arXiv:2210.04620). "
    "[ONC-Cures] Office of the National Coordinator for Health IT, \"Cures Act Final Rule,\" "
    "healthit.gov/regulations/cures-act-final-rule; Federal Register, 21st Century Cures Act: "
    "Interoperability, Information Blocking, and the ONC Health IT Certification Program, 2020. "
    "[TEFCA] ONC Recognized Coordinating Entity (The Sequoia Project), rce.sequoiaproject.org/tefca. "
    "[EHDS25] Regulation (EU) 2025/327 of the European Parliament and of the Council on the European "
    "Health Data Space, Official Journal of the European Union, 5 March 2025.",
    italic=True, size=9, color=MUTED)

page_break(doc)

# ==========================================================================
# SECTION 4 -- THE DIFFICULTY OF WORKING WITH REAL CLINICAL DATA
# ==========================================================================
add_heading(doc, "4. The Difficulty of Working With Real Clinical Data", level=1)
add_para(doc,
    "A technical report on cross-site heterogeneity can make the whole problem sound like a pure "
    "modeling exercise. It is not. The reason so much clinical-ML research relies on the same handful "
    "of datasets, and the reason multi-institution studies are rarer than single-site ones, is a set "
    "of real, well-justified legal and ethical barriers that sit upstream of any modeling choice. "
    "Understanding them is necessary to understand why this workshop's dataset (Section 5) is unusual, "
    "and why federated learning (Section 3.2) is often not a stylistic preference but the only "
    "architecture actually available.")

add_heading(doc, "4.1 IRB review and the Common Rule", level=2)
add_para(doc,
    "In the United States, essentially any use of data involving human subjects for research -- "
    "including a purely retrospective, secondary analysis of existing clinical records -- falls under "
    "the Common Rule (45 CFR 46) and requires review by an Institutional Review Board (IRB) [45CFR46]. "
    "Even when a researcher believes an analysis poses minimal risk and may qualify for an exemption, "
    "that determination itself has to be made by the IRB, not assumed by the researcher. This review "
    "exists for good reason -- it is the primary safeguard for patient rights and welfare in research "
    "-- but it also means that before a single row of data can be analyzed, a formal, often multi-week "
    "institutional process has to run its course, and it has to run separately at every participating "
    "institution in a multi-site study.")

add_heading(doc, "4.2 HIPAA, de-identification, and data use agreements", level=2)
add_para(doc,
    "Separately from IRB review, the HIPAA Privacy Rule governs any data that qualifies as Protected "
    "Health Information (PHI) held by a covered entity. Sharing PHI outside the originating "
    "institution requires patient authorization, a waiver, or de-identification meeting one of two "
    "recognized standards: Safe Harbor (removing eighteen specific categories of identifier) or "
    "Expert Determination (a qualified statistician formally certifies that re-identification risk is "
    "very small) [HIPAA-Deident]. Both routes require real expertise and institutional sign-off, not "
    "just deleting a name column. Moving data between institutions at all typically also requires a "
    "negotiated Data Use Agreement (DUA) -- a legal contract specifying permitted uses, security "
    "controls, and publication rights -- and DUAs are usually negotiated per pair of institutions, "
    "which is a large part of why genuinely multi-site datasets are so much rarer than single-site "
    "ones: a four-hospital study does not need one DUA, it potentially needs up to six.")

add_heading(doc, "4.3 Even \"open\" datasets are usually credentialed, not public", level=2)
add_para(doc,
    "Some of the best-known multi-institution clinical datasets are still gated behind a "
    "credentialing process rather than being truly public downloads. PhysioNet's MIMIC-IV, for "
    "example, requires a signed PhysioNet Credentialed Health Data Use Agreement and completion of "
    "human-subjects-research training (the CITI \"Data or Specimens Only Research\" course) before an "
    "individual researcher can download the data [PhysioNet-MIMIC] -- a reasonable and proportionate "
    "safeguard, but a real barrier for a live, drop-in teaching setting where attendees cannot be "
    "expected to complete a multi-day credentialing process in advance.")

add_table_caption(doc, "Table 4.1 -- The layered barriers between a clinical dataset and a public download")
add_table(doc,
    ["Barrier", "What it requires", "Typical effect"],
    [
        ["IRB / Common Rule review", "Review (or a formal exemption determination) for essentially "
         "any human-subjects research use, including retrospective analysis", "Weeks to months, "
         "separately at every participating institution"],
        ["HIPAA Privacy Rule", "De-identification via Safe Harbor (18 identifiers removed) or Expert "
         "Determination (statistician-certified) before data can leave the covered entity",
         "Requires dedicated expertise; incomplete de-identification blocks release entirely"],
        ["Data Use Agreement (DUA)", "A negotiated legal contract between every pair of sharing "
         "institutions, covering permitted uses, security and publication rights", "Often months per "
         "institution pair; scales roughly quadratically with the number of sites"],
        ["Credentialing (e.g. PhysioNet)", "Signed DUA plus completed human-subjects training (e.g. "
         "CITI) before an individual researcher may download", "Days to weeks per researcher, "
         "periodically renewed"],
        ["Net effect", "The large majority of clinical data never becomes a public, redistributable "
         "research dataset", "Federated learning becomes a practical necessity for genuine multi-site "
         "study, not just a preferred architecture"],
    ])

add_para(doc,
    "This is what makes the dataset in Section 5 genuinely unusual. It was donated to the public UCI "
    "Machine Learning Repository in 1988-89 -- more than a decade before the HIPAA Privacy Rule was "
    "published (December 2000) and before covered entities were required to comply with it (April "
    "2003) [HIPAA-Dates] -- and released with the explicit intent of open reuse. A comparable "
    "four-hospital dataset collected today would, even with every institution fully willing to share, "
    "face a materially higher bar to public release than this one did. That historical accident is "
    "precisely why a workshop can hand attendees real, multi-institution patient data with zero "
    "credentialing, and it is the direct bridge into the next section.")

add_para(doc,
    "Sources for this section: "
    "[45CFR46] U.S. Dept. of Health and Human Services, Office for Human Research Protections, "
    "\"Federal Policy for the Protection of Human Subjects (45 CFR 46),\" hhs.gov/ohrp. "
    "[HIPAA-Deident] U.S. HHS Office for Civil Rights, \"Guidance Regarding Methods for "
    "De-identification of Protected Health Information in Accordance with the HIPAA Privacy Rule,\" "
    "hhs.gov/hipaa/for-professionals/special-topics/de-identification. "
    "[HIPAA-Dates] U.S. HHS, Standards for Privacy of Individually Identifiable Health Information, "
    "Federal Register, 28 December 2000 (final rule); general compliance date 14 April 2003. "
    "[PhysioNet-MIMIC] PhysioNet, \"MIMIC-IV\" data-access requirements (PhysioNet Credentialed "
    "Health Data Use Agreement 1.5.0 and CITI \"Data or Specimens Only Research\" training), "
    "physionet.org/content/mimiciv.",
    italic=True, size=9, color=MUTED)

page_break(doc)

# ==========================================================================
# SECTION 5 -- OUR DATASET
# ==========================================================================
add_heading(doc, "5. Our Dataset: The UCI Heart Disease Federation", level=1)
add_para(doc,
    "A good demonstration dataset for \"interoperability is not comparability\" needs one property "
    "above all else: the four sources must genuinely be interoperable -- same variables, same coding, "
    "same case-report form -- while still differing underneath in ways that matter. A synthetic split "
    "of one hospital's data into four random pieces would be trivially comparable (there is no real "
    "difference to find) and would teach nothing. The UCI Heart Disease federation is close to the "
    "opposite extreme: it is one of the very few open clinical datasets where four *real, independent* "
    "hospitals filled out the *same* form, so every difference we measure is a genuine site effect, "
    "not an artefact of how the demonstration was constructed -- and, as Section 4 explained, one of "
    "the very few such datasets that carries no access barrier at all.")

add_heading(doc, "5.1 The four sites at a glance", level=2)
add_table_caption(doc, "Table 5.1 -- The four hospital sites")
add_table(doc,
    ["Site", "Institution", "Patients", "Any-disease prevalence", "Most common severity"],
    [
        ["Cleveland", "Cleveland Clinic Foundation", "303", "45.9%", "No disease (54.1%)"],
        ["Hungary", "Hungarian Institute of Cardiology", "294", "36.1%", "No disease (63.9%)"],
        ["Switzerland", "Univ. Hospital Zurich / Basel", "123", "93.5%", "Mild (39.0%)"],
        ["V.A. Long Beach", "V.A. Medical Center", "200", "74.5%", "Mild (28.0%)"],
        ["Total", "-- ", "920", "55.3% pooled", "-- "],
    ])
add_para(doc,
    "Prevalence alone ranges nearly threefold, from 36% to 94% -- this is what the report calls "
    "*true* distributional shift: Switzerland is a tertiary referral centre that sees a sicker, "
    "pre-selected population, not a data problem.")

add_heading(doc, "5.2 Two different heterogeneity fingerprints, side by side", level=2)
add_para(doc,
    "The dataset is unusually good at teaching the difference between real population variation and "
    "broken measurement, because it contains a clean example of each, at very different scales.")
add_table_caption(doc, "Table 5.2 -- Two heterogeneity fingerprints, by site")
add_table(doc,
    ["Feature", "What it measures", "Cleveland", "Hungary", "Switzerland", "V.A.", "Fingerprint"],
    [
        ["chol", "Serum cholesterol coded 0 when not recorded", "0.0%", "0.0%", "100.0%", "24.5%",
         "Measurement artefact (site-specific)"],
        ["ca", "# vessels by fluoroscopy, coded '?'/'-9' when missing", "1.3%", "98.6%", "95.9%",
         "99.0%", "Measurement artefact (near-universal outside Cleveland)"],
        ["thal", "Thallium stress-test result, same missing coding", "0.7%", "90.5%", "42.3%",
         "83.0%", "Measurement artefact (near-universal outside Cleveland)"],
        ["target severity", "0-4 disease severity", "mixed", "skews 0", "skews 1-2", "mixed",
         "True distributional shift"],
    ])
add_para(doc,
    "Cholesterol is the textbook example used throughout the workshop framing: Switzerland's field "
    "is not \"low cholesterol,\" it is \"never measured,\" coded as a literal zero directly in the "
    "source file. But the far larger gap, and one this project's pipeline surfaces that a cursory "
    "look would not, is `ca` and `thal`: outside Cleveland, both are missing for 80-99% of patients "
    "at every hospital. Any model trained naively across all four sites is, for those two features, "
    "running on an imputed placeholder value for the overwhelming majority of non-Cleveland patients "
    "-- a measurement gap roughly twenty times larger than the well-known cholesterol example, and "
    "invisible unless someone explicitly checks for it. That is precisely the curation risk this "
    "toolkit is built to catch before it reaches a model.")

add_image(doc, "02_missingness.png",
    "Figure 5.1 -- Fraction of unrecorded values per feature per site, both fingerprints side by "
    "side. Cholesterol's 100% gap at Switzerland is visible; so is the much larger, near-universal "
    "gap in `ca` and `thal` outside Cleveland. Produced by viz.plot_missingness, from "
    "heterogeneity.missingness_report.")

add_para(doc,
    "The heatmap above summarizes the fraction missing; the next figure shows what the measured "
    "values actually look like once the unrecorded cases are excluded rather than silently treated "
    "as zero.")
add_image(doc, "03_chol_distributions.png",
    "Figure 5.2 -- Per-site distribution of measured cholesterol values (Switzerland has none to "
    "plot -- it is marked as fully unrecorded rather than silently included as zero). Produced by "
    "viz.plot_feature_distributions.")

add_heading(doc, "5.3 True label shift: what we predict differs sharply by site", level=2)
add_table_caption(doc, "Table 5.3 -- Severity class counts by site")
add_table(doc,
    ["Severity class", "Cleveland", "Hungary", "Switzerland", "V.A.", "Pooled total"],
    [
        ["0 -- No disease", "164", "188", "8", "51", "411 (44.7%)"],
        ["1 -- Mild", "55", "37", "48", "56", "196 (21.3%)"],
        ["2 -- Moderate", "36", "26", "32", "41", "135 (14.7%)"],
        ["3 -- Severe", "35", "28", "30", "42", "135 (14.7%)"],
        ["4 -- Critical", "13", "15", "5", "10", "43 (4.7%)"],
    ])
add_para(doc,
    "Hungary skews heavily toward \"no disease\"; Switzerland is almost entirely disease-positive "
    "and, unusually, has more \"mild\" cases than \"no disease\" cases -- consistent with a referral "
    "population rather than a screening population.")

add_image(doc, "eda/e01_target_distribution.png",
    "Figure 5.3 -- Pooled 5-class severity distribution, and the derived binary any-disease summary. "
    "Produced by eda.plot_target_distribution.")

add_para(doc,
    "The pooled breakdown above hides which hospital contributes which severities; splitting it out "
    "by site makes the label shift from Table 5.3 directly visible.")
add_image(doc, "eda/e02_target_by_site.png",
    "Figure 5.4 -- The same severity breakdown, split by site -- the label-shift story from Table "
    "5.3 made visual. Produced by eda.plot_target_by_site.")

add_heading(doc, "5.4 Covariate shift: the sites are statistically -- and visibly -- separable", level=2)
add_para(doc,
    "A more rigorous check than eyeballing histograms is to ask: can a simple classifier, given only "
    "a patient's clinical features (never told which hospital they came from), guess the hospital? "
    "If it cannot do better than a coin flip (AUC around 0.5), the sites are statistically "
    "indistinguishable and any model or calibration should transfer cleanly. If it can guess almost "
    "perfectly (AUC near 1.0), the sites occupy different regions of feature space and nothing "
    "trained at one is guaranteed to behave sensibly at another.")
add_table_caption(doc, "Table 5.4 -- Pairwise domain-classifier AUC (5-fold cross-validated logistic regression)")
add_table(doc,
    ["", "Cleveland", "Hungary", "Switzerland", "V.A."],
    [
        ["Cleveland", "0.50", "0.89", "1.00", "0.88"],
        ["Hungary", "0.89", "0.50", "1.00", "0.93"],
        ["Switzerland", "1.00", "1.00", "0.50", "0.88"],
        ["V.A.", "0.88", "0.93", "0.88", "0.50"],
    ])
add_para(doc, "Every off-diagonal value is 0.88 or higher -- these hospitals are almost perfectly "
    "separable from features alone.", italic=True, color=MUTED, size=10)

add_image(doc, "05_domain_auc.png",
    "Figure 5.5 -- The AUC matrix as a heatmap. Produced by viz.plot_divergence_matrix, from "
    "heterogeneity.domain_auc_matrix.")

add_para(doc,
    "The AUC matrix asks whether a classifier can tell sites apart; Jensen-Shannon divergence asks a "
    "narrower, complementary question directly about one feature's distribution shape, without "
    "needing to train anything.")
add_image(doc, "04_js_divergence.png",
    "Figure 5.6 -- Jensen-Shannon divergence between sites for maximum heart rate (thalach): a "
    "second, distribution-shape-based confirmation of the same covariate-shift story, in a [0, 1] "
    "score where 0 means identical distributions. Produced by viz.plot_divergence_matrix, from "
    "heterogeneity.js_divergence_matrix.")

add_para(doc,
    "Having established that sites differ (Figures 5.5-5.6), the next four figures return to the "
    "underlying input data itself and ask which individual features actually drive that difference "
    "and which ones drive the disease-severity prediction task.")
add_image(doc, "eda/e03_continuous_grid.png",
    "Figure 5.7 -- All five continuous features, split by severity class. Where the five severity "
    "curves separate, that feature is informative for the model; where they overlap, it is not. "
    "Produced by eda.plot_continuous_grid.")

add_para(doc,
    "Max heart rate stood out in Figure 5.7 as one of the more separated features by severity; the "
    "box plots below isolate it by site instead, for a cleaner, single-feature read on how its "
    "distribution shifts across hospitals.")
add_image(doc, "eda/e04_thalach_by_site.png",
    "Figure 5.8 -- Max heart rate distribution by site as box plots. Produced by "
    "eda.plot_feature_boxplots_by_site.")

add_para(doc,
    "The continuous features above are only half the picture; the categorical and ordinal clinical "
    "variables (chest-pain type, exercise angina, ST slope, thalassemia result, and others) carry at "
    "least as much signal.")
add_image(doc, "eda/e05_categorical_grid.png",
    "Figure 5.9 -- Mean severity within each level of six categorical features (chest-pain type, "
    "sex, exercise angina, ST slope, resting ECG, thalassemia result). Asymptomatic chest pain, "
    "exercise-induced angina, a flat/down ST slope and a reversible thalassemia defect all carry "
    "substantially higher mean severity -- clinically sensible relationships the model can exploit, "
    "and a sanity check that the loaded data behaves the way cardiology domain knowledge predicts. "
    "Produced by eda.plot_categorical_grid.")

add_para(doc,
    "Finally, viewing every feature relationship at once as a correlation matrix confirms nothing in "
    "Figures 5.7-5.9 is contradicted by the pairwise structure of the data.")
add_image(doc, "eda/e06_correlation.png",
    "Figure 5.10 -- Pearson correlation among all 13 features and the 5-class severity target. "
    "Chest-pain type, exercise angina, ST depression, max heart rate and vessel count are the "
    "strongest correlates of severity. Produced by eda.plot_correlation_heatmap.")

add_heading(doc, "5.5 The secondary task confirms this is not a fluke of one label choice", level=2)
add_para(doc,
    "If the heterogeneity story depended on the specific choice of \"predict disease severity,\" it "
    "would be a weaker argument. The toolkit supports a second, independent task on the same four "
    "sites -- predicting chest-pain type instead of disease severity -- and label shift shows up "
    "there too, just as clearly, confirming the effect is a property of the sites, not of one "
    "particular target.")

add_image(doc, "cp/c01_class_distribution.png",
    "Figure 5.11 -- Chest-pain-type distribution by site, the secondary-task analogue of Figure 5.4. "
    "Hungary is heavy on atypical angina; Switzerland is roughly 80% asymptomatic. Produced by "
    "viz.plot_class_distribution_by_site, via scripts/run_chest_pain_pipeline.py.")

add_heading(doc, "5.6 Summary: what makes this dataset ideal", level=2)
add_bullets(doc, [
    "Genuinely interoperable: all four sites used the same case-report form and the same 13-variable "
    "schema -- there is no format barrier to integration, which isolates the *meaning* problem from "
    "the *plumbing* problem.",
    "Genuinely heterogeneous, and in two distinguishable ways at once: a true population difference "
    "(prevalence 36%-94%) and a measurement artefact (chol, and far more severely, ca/thal), so the "
    "toolkit's core distinction is not hypothetical.",
    "Small enough to train and calibrate live in a workshop (920 patients, sub-second training), "
    "large enough that the effects are statistically real rather than noise (Beta-band-checked in "
    "Section 7.3).",
    "Open, credential-free, and citable, for the historical reasons explained in Section 4 -- no "
    "data-use agreement stands between a workshop attendee and running the code themselves.",
    "The finding replicates on a second, independently chosen prediction task on the same sites "
    "(Figure 5.11), which is evidence the heterogeneity is a property of the hospitals, not an "
    "artefact of the modeling choice.",
])

page_break(doc)

# ==========================================================================
# SECTION 6 -- METHODS OF INTEROPERABILITY
# ==========================================================================
add_heading(doc, "6. Methods of Interoperability: Choosing a Standard", level=1)
add_para(doc,
    "Section 5 relied on a lucky historical accident: four hospitals happened to fill out the same "
    "paper case-report form. A real curation pipeline built today does not get to rely on luck -- it "
    "has to deliberately choose a data standard. This section surveys the realistic options and lays "
    "out exactly what adopting any of them does, and does not, solve.")

add_heading(doc, "6.1 Interoperability has levels, and most mandates stop at the easy ones", level=2)
add_para(doc,
    "HIMSS (the Healthcare Information and Management Systems Society) distinguishes four levels of "
    "interoperability: foundational (data can move from one system to another at all), structural "
    "(the format and structure of the exchanged data is standardized, e.g. a consistent field for "
    "\"cholesterol\"), semantic (the exchanged value carries a standardized, shared meaning, e.g. a "
    "specific LOINC code that both systems interpret identically), and organizational (the "
    "governance, policy and workflow context needed for the exchange to actually happen in practice). "
    "Most current regulatory mandates (Section 3.3) are strongest at the foundational and structural "
    "levels. Semantic interoperability -- the level where this report's entire heterogeneity story "
    "lives -- is the hardest to mandate and the easiest to get only partway to.")

add_heading(doc, "6.2 The standards landscape", level=2)
add_table_caption(doc, "Table 6.1 -- Clinical data standards compared for this workshop's use case")
add_table(doc,
    ["Standard", "Primary use case", "Structural interoperability", "Semantic interoperability", "Fit for multi-site research"],
    [
        ["HL7 v2 messaging", "Legacy point-to-point clinical messaging (labs, admissions, orders)",
         "Moderate -- widely implemented but loosely enforced", "Weak -- minimal standard "
         "vocabulary binding", "Poor: still the production backbone in many hospitals, but not "
         "designed for pooled research analysis"],
        ["HL7 FHIR", "Modern REST/JSON API exchange; the standard named in current US/EU "
         "interoperability mandates", "Strong", "Moderate-to-strong when paired with standard "
         "terminologies (LOINC, SNOMED CT, RxNorm) -- inconsistently adopted site to site",
         "Good for live exchange; not itself a research-analysis data model"],
        ["OMOP Common Data Model (OHDSI)", "Standardizing observational health data from many "
         "institutions into one schema and vocabulary for pooled/federated research",
         "Strong", "Strong -- the standard vocabulary mapping is the point of the model",
         "Excellent: purpose-built for exactly this report's use case"],
        ["openEHR", "Archetype-based clinical modeling with very high semantic expressiveness",
         "Strong", "Strong", "Good in principle; limited adoption in US hospital systems specifically"],
        ["CDISC SDTM / ADaM", "Regulated clinical-trial data submission to FDA/EMA", "Strong",
         "Strong, within the trial-data domain", "Poor fit: built for trials, not observational "
         "multi-site EHR data"],
        ["DICOM", "Medical imaging exchange and storage", "Strong (imaging-specific)",
         "Strong (imaging-specific)", "Not applicable to this workshop's tabular clinical variables, "
         "but the correct standard once imaging enters a curation pipeline"],
    ])

add_heading(doc, "6.3 What standardization can and cannot buy you", level=2)
add_para(doc,
    "Table 6.1 is a starting point for that choice, not a verdict -- which standard (or combination) "
    "fits a given curation pipeline depends on institutional context, existing vendor systems, and "
    "whether the goal is live clinical exchange or pooled research analysis, and is a decision this "
    "report leaves to the team building the pipeline rather than asserting on their behalf.")
add_para(doc,
    "What the report *can* say with confidence is the limit of what any of these standards buys you. "
    "Adopting a common schema and a common vocabulary reduces, but does not eliminate, "
    "measurement-process heterogeneity -- a hospital can code cholesterol to the correct standard "
    "concept and still simply order that test less often than another hospital does, which reproduces "
    "Switzerland's missingness pattern (Section 5.2) inside a perfectly standards-compliant pipeline. "
    "No coding standard changes the fact that a tertiary referral centre sees a sicker, pre-selected "
    "population either. That is why this toolkit's heterogeneity diagnostics (Section 5) and "
    "conformal-coverage diagnostics (Sections 7-8) belong *downstream* of any standardization effort "
    "as a verification step, not as a one-time substitute for one.")
add_para(doc,
    "It is worth being explicit that the UCI dataset itself predates and does not use any of the "
    "standards in Table 6.1 -- it is distributed as raw, fixed-format, per-site flat files. That is a "
    "feature for this report's purposes, not a flaw: it shows what pre-standardization interoperability "
    "looked like, and every figure in Section 5 is a direct, honest picture of what adopting a modern "
    "standard would, and would not, have caught. Standardizing the format would have caught nothing "
    "about the true prevalence swing across sites (Table 5.1) -- that is real biology, not a coding "
    "problem -- while a well-governed pipeline that tracks per-site data completeness by design would "
    "plausibly have surfaced the `ca`/`thal` measurement gap (Table 5.2) earlier than this report did.",
    italic=True, color=MUTED)

page_break(doc)

# ==========================================================================
# SECTION 7 -- WHY CONFORMAL PREDICTION HELPS
# ==========================================================================
add_heading(doc, "7. Why Conformal Prediction Helps With Site-Level Heterogeneity", level=1)

add_heading(doc, "7.1 The problem: models are confident even when they are wrong", level=2)
add_para(doc,
    "A typical machine-learning classifier reports a number that looks like a probability -- "
    "\"92% chance this patient has heart disease.\" That number is not actually a probability in "
    "any guaranteed sense; it is whatever the model happened to output. Nothing forces it to be "
    "correct 92% of the time. In practice, a model trained on one hospital's patients and then run "
    "on a different hospital's patients often stays just as confident -- the 92% does not shrink -- "
    "while its actual correctness rate quietly drops. The model does not know it has left home. This "
    "is the single most dangerous failure mode in cross-institution deployment: not a model that "
    "visibly fails, but one that fails silently while sounding just as sure of itself as it did "
    "during development.")
add_para(doc,
    "Interoperability standards (Section 6) solve the problem of *getting the data there*. They do "
    "nothing to solve this problem, because the failure is not about the data format -- it is about "
    "whether the statistical relationship the model learned at Hospital A still holds at Hospital B. "
    "Two hospitals can use an identical schema, an identical column called \"cholesterol,\" and still "
    "have that column mean something different in practice (measured routinely at one site, almost "
    "never measured at another). A model cannot tell the difference between a hospital with "
    "genuinely different patients and a hospital with a broken measurement pipeline unless something "
    "is specifically built to check.")

add_heading(doc, "7.2 What conformal prediction actually does", level=2)
add_para(doc,
    "Conformal prediction is a wrapper you can put around *any* trained model -- logistic "
    "regression, a neural network, anything that outputs a score per class -- that converts its "
    "raw scores into a set of plausible answers, with a guarantee attached to the size and content "
    "of that set. Instead of the model saying \"heart disease, severity 2, 92% confident,\" a "
    "conformal predictor says \"the true severity is somewhere in {1, 2, 3}, and this kind of "
    "statement is right at least 90% of the time.\" When the model is unsure, the set grows to "
    "include more possibilities; when it is sure, the set can shrink to a single class. Crucially, "
    "the 90% figure is not a hope -- it is a mathematical guarantee that holds regardless of whether "
    "the underlying model is any good, as long as one condition holds: the data you calibrate the "
    "guarantee on must look statistically like the data you later apply it to.")
add_para(doc,
    "That condition -- formally called *exchangeability*, discussed in Section 8 -- is exactly where "
    "site heterogeneity enters the picture. Calibrating a conformal predictor at one hospital and "
    "deploying it at another silently assumes the two hospitals are exchangeable. If they are not, "
    "the 90% promise degrades to whatever the real cross-site agreement supports, and conformal "
    "prediction becomes a *measuring instrument* for that degradation: you calibrate at one site, "
    "check the empirical coverage at every other site, and the gap between the promised rate and the "
    "observed rate is a direct, quantitative readout of how much that pair of sites disagrees. No "
    "other common technique gives you a single interpretable number for \"how much can I trust this "
    "model somewhere it hasn't been tested.\"")

add_heading(doc, "7.3 Averages hide the failure -- you have to look per site", level=2)
add_para(doc,
    "The single most important lesson in this toolkit is that a pooled, dataset-wide coverage number "
    "can look perfect while hiding a badly broken subgroup. A conformal predictor calibrated across "
    "all four hospitals together can report 90% coverage overall while one hospital sits at 70% and "
    "the others compensate by over-covering. This is exactly what we observe: a model federated over "
    "Cleveland, Hungary and the V.A., then calibrated using only Cleveland patients and deployed "
    "everywhere, hits 90-91% coverage at Cleveland, Hungary and the V.A. -- and drops to 71% at "
    "Switzerland, the hospital the federation never trained on (Figure 7.3). Reported as one pooled "
    "number, this problem is invisible. Reported per site, it is impossible to miss. That is the "
    "practical argument for conformal prediction in a curation pipeline: it is cheap to compute, "
    "attaches to any model you already have, and turns \"is this data comparable across sites\" from "
    "a qualitative judgment call into a number you can put in a table and set a threshold on.")

add_image(doc, "01_site_overview.png",
    "Figure 7.1 -- Patients per site and disease prevalence per site. Prevalence swings from "
    "36% (Hungary) to 94% (Switzerland): a genuine difference in patient population (true "
    "distributional shift), not a data-entry artefact. Produced by viz.plot_site_overview, called "
    "from data.summarize_sites.")

add_para(doc,
    "Before showing the empirical result, it helps to see the distinction it depends on drawn "
    "schematically first.")
add_image(doc, "paper/fig10_coverage_notions.png",
    "Figure 7.2 -- The distinction this whole report rests on, drawn schematically for two groups. "
    "Left: no coverage guarantee at all. Middle: 'marginal' coverage -- the overall average looks "
    "fine (green dots) while one group is quietly failing (concentrated red dots). Right: "
    "'conditional' coverage -- every group is covered at the target rate individually. Produced by "
    "paper_figures.fig10_coverage_notions, a recreation of Figure 10 in Angelopoulos & Bates (2022).")

add_para(doc,
    "That schematic is not hypothetical here -- this is the real measurement on this project's data.")
add_image(doc, "09_coverage_by_site.png",
    "Figure 7.3 -- The headline empirical result. A conformal predictor calibrated only on "
    "Cleveland data is deployed at all four hospitals. It meets the 90% target at Cleveland, Hungary "
    "and the V.A., and silently drops to 71% at Switzerland -- well outside the shaded 'benign "
    "fluctuation' band that captures ordinary sampling noise, meaning this drop is a real effect, "
    "not chance. Produced by viz.plot_coverage_by_site, using conformal.LACPredictor and "
    "evaluate.coverage.")

add_para(doc,
    "Figure 7.3 fixed the calibration site at Cleveland; the transfer matrix below generalizes that "
    "to every possible calibrate/deploy pairing at once.")
add_image(doc, "08_transfer_matrix.png",
    "Figure 7.4 -- The full picture: calibrate at the row's hospital, deploy at the column's "
    "hospital, read off coverage. The diagonal (calibrate and deploy at the same hospital) is "
    "healthy everywhere. Off-diagonal cells show what happens when a hospital's calibration is "
    "exported elsewhere -- deploying anything calibrated on Cleveland, Hungary or the V.A. at "
    "Switzerland lands around 70-72%, while a Switzerland-calibrated predictor over-covers "
    "everywhere else (its patients are on average sicker, so its threshold is looser than it needs "
    "to be elsewhere). Produced by viz.plot_transfer_matrix.")

add_table_caption(doc, "Table 7.1 -- Coverage of a Cleveland-calibrated predictor, by deployment site "
    "(target: 90%, calibration set = 303 Cleveland patients, split-conformal LAC score, α = 0.10)")
add_table(doc,
    ["Deployment site", "Empirical coverage", "Gap vs. 90% target", "Verdict"],
    [
        ["Cleveland (same site as calibration)", "90.8%", "+0.8 pp", "Meets target"],
        ["Hungary", "90.8%", "+0.8 pp", "Meets target"],
        ["V.A. Long Beach", "90.5%", "+0.5 pp", "Meets target"],
        ["Switzerland (never seen by the federated model)", "70.7%", "−19.3 pp", "Fails, silently"],
    ])
add_para(doc,
    "This is the practical payoff for a curation pipeline: instead of trusting that a model \"should\" "
    "generalize because the incoming data is schema-compliant, you get a per-site number that tells "
    "you exactly where trust breaks down, before a clinician or downstream system relies on it.",
    italic=True, color=MUTED)

page_break(doc)

# ==========================================================================
# SECTION 8 -- THE MATHEMATICS
# ==========================================================================
add_heading(doc, "8. The Mathematics, and What Every Piece of Code Produces", level=1)
add_para(doc,
    "This section explains the method in plain language first, gives the exact formula each idea "
    "corresponds to, and then says precisely which file and which function in the toolkit implements "
    "it and which figure it produces. Everything here is implemented in plain NumPy -- no black-box "
    "library -- specifically so every step can be read line by line.")
add_para(doc,
    "Before any of that, it helps to see the destination first -- what a prediction set actually "
    "looks like -- exactly the role Figure 1 plays at the start of the source paper, before its own "
    "Section 2 explains the method formally.")

add_image(doc, "paper/fig01_prediction_set_examples.png",
    "Figure 8.1 -- Three real patients from this project's own data, all truly \"No disease,\" and "
    "the prediction set the calibrated model actually returns for each. Unlike every other figure in "
    "this section, this one is not illustrative: it trains the real federated model, calibrates a "
    "real APS conformal predictor (alpha = 0.10), and searches the held-out test patients for the "
    "most confident correctly-covered singleton, a representative 3-class case, and the least "
    "confident correctly-covered 4-class case -- all restricted to patients whose true label is "
    "\"No disease,\" the only one of the five severity classes for which a full 1-4 spread of "
    "covered set sizes exists. In the hardest case the model's own top guess is \"Severe,\" not the "
    "true \"No disease\" -- the true class is fourth, not first -- yet it still survives inside the "
    "set, which is the entire point of the coverage guarantee: it protects the true answer even when "
    "the model itself is confidently pointing elsewhere. This is a direct, non-illustrative echo of "
    "the source paper's own Figure 1, where a third fox-squirrel photograph fools the classifier into "
    "ranking \"marmot\" first. Produced by paper_figures.fig01_prediction_set_examples, via "
    "scripts/run_paper_figure_recreations.py.")

add_heading(doc, "8.1 The one assumption everything rests on: exchangeability", level=2)
add_para(doc,
    "Conformal prediction needs exactly one statistical assumption, and it is weaker than almost "
    "anything else in statistics: the calibration data and the test data must be *exchangeable* -- "
    "informally, drawn from the same underlying process, so that shuffling the order of the combined "
    "set would not tell you anything about which points were 'calibration' and which were 'test'. It "
    "does not require the data to be Gaussian, linearly separable, or generated by any particular "
    "model at all. This is what makes conformal prediction so widely applicable -- and it is also "
    "exactly the assumption that cross-hospital deployment violates. Calibration data from Cleveland "
    "and test data from Switzerland are not exchangeable if Switzerland patients are drawn from a "
    "different population (94% disease prevalence vs. Cleveland's 46%). The guarantee below is "
    "conditional on an assumption that this dataset is specifically built to violate on purpose, so "
    "that the violation is visible.")

add_heading(doc, "8.2 The nonconformity score and the conformal quantile", level=2)
add_para(doc,
    "Step one: define a single number, the 'nonconformity score,' that measures how surprising a "
    "true label was, given what the model predicted. A large score means the model's output looked "
    "nothing like the truth; a small score means the model was basically right. Step two: compute "
    "this score for every point in a held-out 'calibration' set the model never trained on. Step "
    "three: take a specific, slightly-inflated quantile of those calibration scores -- inflated just "
    "enough to make the finite-sample guarantee exact rather than approximate:")
add_math(doc, [
    ("q̂  =  Quantile( s", "normal"), ("1", "sub"), (", …, s", "normal"), ("n", "sub"),
    (" ;  ⌈(n + 1)(1 − α)⌉ / n )", "normal"),
])
add_para(doc,
    "Here n is the number of calibration points and α (alpha) is the error rate you are willing "
    "to accept (α = 0.10 throughout this report, i.e. a 90% target). Step four: for any new "
    "patient, build the prediction set by keeping every candidate label whose score would not have "
    "exceeded q̂ (\"q-hat\"):")
add_math(doc, [("C(x)  =  { y : s(x, y) ≤ q̂ }", "normal")])
add_para(doc,
    "This is the entire method. The proof that P(true label ∈ C(x)) ≥ 1 − α is a short "
    "exchangeability argument (Angelopoulos & Bates, 2022, Section 2) -- the important practical "
    "point is that it works for *any* nonconformity score and *any* underlying model. "
    "Implementation: conformal.conformal_quantile(scores, alpha). Both predictors described below "
    "call it as their calibration step.")

add_heading(doc, "8.3 Two ways to define the score: LAC and APS", level=2)
add_para(doc,
    "\"LAC\" (Least Ambiguous set-valued Classifier) uses the simplest possible score: one minus the "
    "model's softmax probability on the true class.")
add_math(doc, [
    ("s", "normal"), ("i", "sub"), ("  =  1 − f̂(x", "normal"), ("i", "sub"), (")", "normal"),
    ("y", "sub"), ("i", "sub"),
])
add_para(doc,
    "A score near 0 means the model put almost all its probability mass on the correct class; a "
    "score near 1 means the model was confidently wrong. Implementation: conformal.lac_scores and "
    "conformal.LACPredictor.")
add_para(doc,
    "\"APS\" (Adaptive Prediction Sets, Romano, Sesia & Candes, 2020) instead sorts the classes from "
    "most to least likely and adds up probability mass until it reaches the true class:")
add_math(doc, [
    ("s(x, y)  =  ", "normal"), ("Σ", "normal"), ("j=1", "sub"), ("o(y)", "sup"),
    ("  π̂", "normal"), ("(j)", "sub"), ("(x)", "normal"),
])
add_para(doc,
    "where the classes are sorted by decreasing predicted probability π̂"
    "₁(x) ≥ π̂₂(x) ≥ … , π̂"
    "₍ⱼ₎(x) denotes the probability of the class ranked j-th in that order, and o(y) "
    "is the rank of the true class y itself -- so the score is just the cumulative probability mass "
    "swept up to and including the true class.")
add_para(doc,
    "LAC tends to build the smallest possible sets on average; APS tends to size its sets more "
    "sensibly example-by-example -- growing more for genuinely ambiguous patients and shrinking more "
    "for clear-cut ones -- which matters most in the multiclass setting used later in this report. "
    "Implementation: conformal.aps_scores and conformal.APSPredictor.")

add_image(doc, "paper/fig02_conformal_illustration.png",
    "Figure 8.2 -- The method end to end, worked through visually on a 5-class example. Panel 1: "
    "compute the score on one calibration point (how far the model's probability on the true class "
    "fell short of 1). Panel 2: histogram the scores across the whole calibration set and mark the "
    "quantile q̂. Panel 3: for a brand-new patient, keep every class whose predicted probability "
    "clears the 1 - q̂ bar -- the surviving classes are the prediction set. Produced by "
    "paper_figures.fig02_conformal_illustration (Figure 2 of Angelopoulos & Bates, recreated with "
    "the same 5-class severity framing used throughout this report).")

add_para(doc,
    "Figure 8.2 showed the LAC score end to end; the same three-step recipe applies to APS, with the "
    "sorting-and-accumulating mechanic shown separately below.")
add_image(doc, "paper/fig04_adaptive_prediction_sets.png",
    "Figure 8.3 -- The APS mechanic specifically: sort classes by predicted probability (left), "
    "accumulate that probability (right) until the running total crosses q̂, and the classes up "
    "to that point form the set. Produced by paper_figures.fig04_aps_illustration (Figure 4).")

add_para(doc,
    "Both illustrations above use synthetic scores chosen to make the mechanic legible; the histogram "
    "below is the real LAC calibration-score distribution from this project's federated model.")
add_image(doc, "07_calibration_scores.png",
    "Figure 8.4 -- The actual calibration-score histogram from this project's data (not a synthetic "
    "illustration): LAC scores from a model trained across Cleveland, Hungary and the V.A., "
    "calibrated on a held-out pool. The red line is q̂ -- everything to its left calibrates the "
    "90% target. Produced by viz.plot_calibration_scores inside scripts/run_end_to_end_pipeline.py.")

add_heading(doc, "8.4 Why the calibration-set size matters", level=2)
add_para(doc,
    "Because q̂ is estimated from a finite sample, the *realized* coverage on any particular test "
    "set is itself a random quantity -- it will not sit exactly on 90% every time, even when every "
    "assumption holds. Vovk (2012) showed this randomness follows a Beta distribution with parameters "
    "tied directly to the calibration-set size n and α:")
add_math(doc, [
    ("Coverage  ∼  Beta( n + 1 − ℓ , ℓ )     where     ℓ = ⌊ (n + 1) α ⌋", "normal"),
])
add_para(doc,
    "With only 50 calibration points, observed coverage can swing several percentage points around "
    "the target just from sampling luck; with 1,000, it is tight. This matters directly for this "
    "report's headline result: Switzerland's 71% coverage is far outside this benign-fluctuation "
    "band even for a calibration set the size of Cleveland's (303 patients), which is how we know the "
    "drop is a real site effect and not noise, rather than merely asserting it. Implementation: "
    "evaluate.coverage_beta_interval, plotted analytically by viz.plot_coverage_beta.")

add_image(doc, "11_coverage_beta.png",
    "Figure 8.5 -- The Beta distribution of coverage at three calibration-set sizes (50, 150, 1000). "
    "The smaller the calibration set, the wider and more uncertain the true coverage really is around "
    "the 90% target -- context for judging whether an observed drop is a real site effect or just "
    "noise. Produced by viz.plot_coverage_beta.")

add_para(doc,
    "Figure 8.5 used calibration sizes matched to this project's actual sites (50-1,000 patients); "
    "the recreation below repeats the same idea at the much larger scale used in the original paper, "
    "for comparison.")
add_image(doc, "paper/fig11_coverage_distribution.png",
    "Figure 8.6 -- The same idea recreated at the scale used in the original paper (n = 100, 1,000, "
    "10,000), showing how sharply the distribution narrows as the calibration set grows. Produced by "
    "paper_figures.fig11_coverage_distribution (Figure 11).")

add_heading(doc, "8.5 Adaptivity: are the sets actually bigger where the model is unsure?", level=2)
add_para(doc,
    "A trivial way to hit 90% coverage is to always output every possible class -- technically "
    "correct, completely useless. A good conformal predictor instead produces small sets for easy "
    "cases and large sets for hard ones. We check this directly by tabulating the distribution of set "
    "sizes at each hospital: a predictor that is well calibrated but *not* adaptive would show the "
    "same size distribution everywhere; ours visibly widens at sites with more label ambiguity. "
    "Implementation: evaluate.set_sizes / evaluate.average_set_size, plotted by "
    "viz.plot_set_size_distribution.")

add_image(doc, "10_set_sizes.png",
    "Figure 8.7 -- Distribution of prediction-set sizes by site (5-class disease-severity task). "
    "Larger sets at a given site mean the model is systematically less certain about patients from "
    "that hospital. Produced by viz.plot_set_size_distribution.")

add_heading(doc, "8.6 The subtler failure: class-conditional coverage", level=2)
add_para(doc,
    "Marginal coverage (the 90% averaged over every patient) can hide a second, independent kind of "
    "failure: it can hit target overall while badly under-covering one specific class, particularly a "
    "rare one, exactly the way it can hide a badly-covered site. We demonstrate this on the toolkit's "
    "secondary task -- predicting chest-pain type (4 classes) rather than disease severity. Marginal "
    "coverage across the pooled test set sits at 90%, but broken out by the true class, the rarest "
    "category (typical angina, only 23 test patients) is covered just 26% of the time, while the most "
    "common category (asymptomatic, 200 patients) is over-covered at 95.5%. Implementation: "
    "evaluate.class_conditional_coverage, plotted by viz.plot_class_conditional_coverage.")

add_image(doc, "cp/c04_class_conditional.png",
    "Figure 8.8 -- Coverage broken out by true chest-pain class. The rare class (typical angina, "
    "n = 23) is covered only 26% of the time against a 90% target -- invisible in the pooled 90% "
    "marginal number, and exactly the sub-population a curation pipeline is most likely to serve "
    "under-represented patients in. Produced by viz.plot_class_conditional_coverage inside "
    "scripts/run_chest_pain_pipeline.py.")

add_heading(doc, "8.7 Beyond classification: the same idea generalizes", level=2)
add_para(doc,
    "The prediction-set idea is not specific to classification. paper_figures.py recreates three "
    "further illustrations from Angelopoulos & Bates showing the same 'score, quantile, widen' "
    "recipe applied to (a) quantile regression, where the conformal step widens a fitted "
    "prediction interval by exactly q̂ to reach guaranteed coverage; (b) a generic point "
    "prediction plus an uncertainty scalar u(x), widened the same way; and (c) a Bayesian posterior "
    "predictive density, where the prediction set becomes the region above a probability threshold. "
    "These three figures use small synthetic examples, exactly as the original paper does, rather "
    "than this project's clinical data -- they are included for conceptual completeness (the "
    "workshop's method is not limited to classification) and are not part of the empirical results "
    "in Sections 5, 7 or 9.")

add_image(doc, "paper/fig06_conformalized_quantile_regression.png",
    "Figure 8.9 -- Conformalized quantile regression: a fitted quantile band (dashed) widened by "
    "q̂ (solid) to reach the coverage guarantee. Synthetic illustration. Produced by "
    "paper_figures.fig06_cqr_illustration (Figure 6).")

add_para(doc,
    "The quantile-regression band above widens an already-interval-valued prediction; the next "
    "illustration instead starts from a single point prediction and conformalizes it with a "
    "generic uncertainty scalar.")
add_image(doc, "paper/fig08_uncertainty_scalar.png",
    "Figure 8.10 -- A point prediction f(x) with a symmetric conformalized band q̂ · u(x). "
    "Synthetic illustration. Produced by paper_figures.fig08_uncertainty_scalar (Figure 8).")

add_para(doc,
    "The final variant swaps a frequentist point prediction for a full Bayesian posterior predictive "
    "density, and takes the prediction set to be the region where that density is highest.")
add_image(doc, "paper/fig09_conformalized_bayes.png",
    "Figure 8.11 -- Conformalized Bayes: the prediction set as the superlevel set of a posterior "
    "predictive density. Synthetic illustration. Produced by paper_figures.fig09_bayes_superlevel "
    "(Figure 9).")

add_heading(doc, "8.8 Code-to-purpose map", level=2)
add_table_caption(doc, "Table 8.1 -- Every module in src/fedconformal/, its responsibility, and the "
    "figures it produces")
add_table(doc,
    ["Module", "Responsibility", "Key functions", "Figures produced"],
    [
        ["data.py", "Load the 4 raw UCI site files, impute missing values, standardize, split "
                    "into per-site arrays for either task", "load_raw, load_sites, preprocess, "
                    "summarize_sites", "(feeds every figure; produces no plots itself)"],
        ["heterogeneity.py", "Quantify site differences before any modeling: label shift, "
                    "missingness, distributional divergence, a domain-classifier alarm",
                    "missingness_report, js_divergence_matrix, domain_auc_matrix",
                    "Figs. 5.1, 5.5, 5.6"],
        ["eda.py", "First-look exploratory plots of the raw inputs: target, features, "
                    "correlation structure", "plot_target_distribution, plot_continuous_grid, "
                    "plot_categorical_grid", "Figs. 5.3, 5.4, 5.7-5.10"],
        ["federated.py", "A from-scratch FedAvg simulator (binary logistic or softmax model, "
                    "chosen automatically by class count) plus a centralized baseline",
                    "federated_averaging, train_centralized, make_model", "Figs. 9.1, 9.2"],
        ["conformal.py", "The conformal-prediction core: nonconformity scores, the quantile, "
                    "and the two set-valued predictors", "conformal_quantile, lac_scores, "
                    "aps_scores, LACPredictor, APSPredictor",
                    "Underlies every coverage/set figure (Figs. 7.3, 7.4, 8.2, 8.3, ...)"],
        ["evaluate.py", "Every coverage/size metric used to grade a conformal predictor, plus "
                    "the analytic Beta band for judging noise vs. signal", "coverage, "
                    "average_set_size, size_stratified_coverage, class_conditional_coverage, "
                    "coverage_beta_interval", "Figs. 8.5, 8.7, 8.8"],
        ["viz.py", "All plotting -- one function per figure, a fixed colorblind-safe site "
                    "palette, a shared visual style", "plot_site_overview, "
                    "plot_coverage_by_site, plot_transfer_matrix, plot_class_conditional_coverage, ...",
                    "nearly every figure in this report"],
        ["paper_figures.py", "Recreations of the explanatory diagrams from Angelopoulos & Bates "
                    "(2022), for teaching -- synthetic illustrations, not this project's data",
                    "fig01...fig11", "Figs. 7.2, 8.1, 8.2, 8.3, 8.6, 8.9-8.11"],
    ])

page_break(doc)

# ==========================================================================
# SECTION 9 -- PIPELINE WALKTHROUGH
# ==========================================================================
add_heading(doc, "9. Pipeline Walkthrough and Justification: Is This Enough?", level=1)
add_para(doc,
    "This section traces one execution of the full pipeline end to end, then gives an honest "
    "assessment of what is solid, what is a reasonable simplification for a workshop setting, and "
    "what we would add if this were being extended into a production curation tool rather than a "
    "teaching toolkit.")

add_heading(doc, "9.1 Stage 3: federated training actually converges", level=2)
add_para(doc,
    "Before conformal calibration means anything, the underlying model has to have actually learned "
    "something. FedAvg trains a shared softmax model across Cleveland, Hungary and the V.A. (holding "
    "Switzerland out entirely, to test genuine external validation) by having each site take a few "
    "local gradient steps and averaging the resulting weights, weighted by site size, every "
    "communication round -- no patient-level data ever leaves its hospital.")

add_image(doc, "06_fed_curves.png",
    "Figure 9.1 -- Per-site training loss across FedAvg communication rounds, 5-class disease task. "
    "All three participating sites' losses fall together, confirming the federated procedure "
    "converges. Produced by viz.plot_fed_learning_curves.")

add_para(doc,
    "The same convergence check, repeated on the chest-pain task, confirms FedAvg is not something "
    "that happens to work only for the disease-severity target.")
add_image(doc, "cp/c02_fed_curves.png",
    "Figure 9.2 -- The same convergence check on the secondary chest-pain task, confirming FedAvg "
    "behaves consistently across two different targets. Produced by viz.plot_fed_learning_curves, "
    "via scripts/run_chest_pain_pipeline.py.")

add_heading(doc, "9.2 Does the story change with a different task and a different score function?", level=2)
add_para(doc,
    "A natural objection: does the whole story depend on the particular choice of disease-severity "
    "prediction and the LAC score? scripts/compare_prediction_tasks.py runs the identical analysis "
    "-- label shift, covariate shift, cross-site conformal transfer, using the APS score this time -- "
    "on both the 5-class disease task and the 4-class chest-pain task, side by side.")
add_table_caption(doc, "Table 9.1 -- The 5-class disease task and the 4-class chest-pain task, head to head")
add_table(doc,
    ["Metric", "Disease task (5-class)", "Chest-pain task (4-class)"],
    [
        ["Features used", "13 clinical variables", "13 (the other 12 + disease severity)"],
        ["Label-shift (Jensen-Shannon div.)", "0.128", "0.073"],
        ["Covariate-shift (mean domain AUC)", "0.930", "0.929"],
        ["Mean cross-site coverage (target 90%)", "90.9%", "91.0%"],
        ["Worst-case site coverage", "87%", "88%"],
        ["Avg. conformal set size |C(x)|", "2.68 labels", "2.25 labels"],
        ["Set size normalized by class count", "0.536 (53.6% of label space)", "0.563 (56.3%)"],
    ])
add_para(doc,
    "The disease-severity task carries more label shift (0.128 vs. 0.073) -- consistent with the "
    "36%-94% prevalence swing being a bigger effect than chest-pain-type variation -- while both "
    "tasks show essentially identical covariate shift (~0.93 AUC) and both keep cross-site coverage "
    "close to target under APS, with the worst individual site still landing within a few points of "
    "90%. The conclusion that these four hospitals are meaningfully heterogeneous, and that "
    "conformal prediction quantifies the consequence, is not an artefact of one task or one score "
    "function.")

add_image(doc, "compare/task_comparison.png",
    "Figure 9.3 -- The comparison from Table 9.1, plotted: heterogeneity scores, average prediction-"
    "set size, and worst-case cross-site coverage, disease task vs. chest-pain task, side by side. "
    "Produced by compare_prediction_tasks.plot_comparison.")

add_heading(doc, "9.3 The secondary task's own conformal pipeline, run in full", level=2)
add_para(doc,
    "For completeness, the chest-pain task's own cross-site transfer matrix and set-size distribution "
    "-- the same diagnostics as Section 7, applied to a different label -- are included below rather "
    "than only summarized in Table 9.1.")

add_image(doc, "cp/c05_transfer_matrix.png",
    "Figure 9.4 -- Cross-site coverage transfer matrix for the chest-pain task (APS score). Produced "
    "by viz.plot_transfer_matrix, via scripts/run_chest_pain_pipeline.py.")

add_para(doc,
    "The transfer matrix above summarizes every calibrate/deploy pairing at once; fixing the "
    "calibration site at Cleveland specifically -- the same headline framing as Figure 7.3 -- makes "
    "the single-predictor case easier to read directly.")
add_image(doc, "cp/c06_coverage_by_site.png",
    "Figure 9.5 -- Coverage of a Cleveland-calibrated chest-pain predictor across all four sites. "
    "Produced by viz.plot_coverage_by_site, via scripts/run_chest_pain_pipeline.py.")

add_para(doc,
    "Coverage alone does not show adaptivity; the set-size distribution below is the chest-pain "
    "task's analogue of Figure 8.7.")
add_image(doc, "cp/c07_set_sizes.png",
    "Figure 9.6 -- Prediction-set-size distribution by site, chest-pain task. Produced by "
    "viz.plot_set_size_distribution, via scripts/run_chest_pain_pipeline.py.")

add_para(doc,
    "Rounding out the secondary-task diagnostics, the calibration-score histogram below is the "
    "chest-pain analogue of Figure 8.4.")
add_image(doc, "cp/c03_calibration_scores.png",
    "Figure 9.7 -- The chest-pain task's own calibration-score histogram, the secondary-task analogue "
    "of Figure 8.4. Produced by viz.plot_calibration_scores, via scripts/run_chest_pain_pipeline.py.")

add_heading(doc, "9.4 Verdict: is the pipeline right as it stands?", level=2)
add_para(doc,
    "Yes, for its stated purpose. Every claim this report makes is backed by a figure generated from "
    "a real, reproducible run (all 8 notebooks and 6 scripts execute cleanly end to end, and 10 "
    "automated correctness tests -- including a check that split-conformal coverage lands within the "
    "expected Beta band on data with a known-correct answer -- pass). The pipeline correctly "
    "separates the two kinds of heterogeneity the workshop is built to teach, demonstrates the "
    "coverage-guarantee failure concretely and repeatably, and generalizes across two independent "
    "prediction tasks and two different nonconformity scores. As a hands-on teaching and diagnostic "
    "tool for the stated workshop goals, it is complete.")
add_para(doc,
    "That said, extending it toward a production curation tool would benefit from a few additions "
    "that are out of scope for a workshop but worth flagging honestly:")
add_bullets(doc, [
    "A calibration/reliability diagram for the raw model probabilities (expected calibration error), "
    "shown side by side with the conformal coverage story, would make the opening claim -- "
    "\"models are confident even when wrong\" -- directly measurable rather than only illustrated.",
    "Feature-stratified coverage by patient subgroup (sex, age band), not only by site, is already "
    "implemented in evaluate.feature_stratified_coverage but is not currently plotted anywhere -- an "
    "easy addition that would extend the class-conditional-coverage lesson (Section 8.6) to "
    "demographic fairness within a single site.",
    "The per-site mean-imputation of `ca`/`thal` (necessary so the model can train at all, given "
    "80-99% missingness outside Cleveland) is silent by default; adding an explicit 'was this value "
    "missing' indicator feature per patient is a small change that would let the model, and the "
    "curation team, distinguish a genuinely low vessel count from an imputed placeholder.",
    "The mitigation strategies the notebooks currently pose only as end-of-notebook exercises "
    "(per-site / group-balanced calibration, importance-weighted conformal prediction under known "
    "covariate shift) are the natural next experiment -- implementing and plotting at least one of "
    "them would close the loop from \"here is the problem\" to \"here is a first fix.\"",
    "All four sites here are from one country each; a genuinely external validation site (a fifth "
    "hospital never touched during either development or the heterogeneity analysis itself) would "
    "strengthen the claim that the observed coverage drop generalizes beyond this specific "
    "four-hospital federation.",
    "The interoperability-standards survey in Section 6 is descriptive, not implemented -- actually "
    "mapping this dataset onto a common data model such as OMOP, with per-concept completeness "
    "reported the way a real OHDSI network study would, would turn Section 6 from a comparison of "
    "options into a second, standards-based measurement of the same heterogeneity story.",
])
add_para(doc,
    "None of these are required to support the report's central claims -- each is a natural next "
    "increment, not a gap that undermines what is already demonstrated.",
    italic=True, color=MUTED)

page_break(doc)

# ==========================================================================
# REFERENCES
# ==========================================================================
add_heading(doc, "References", level=1)
add_para(doc,
    "Every specific factual claim in this report that is not this project's own reproducible output "
    "(Sections 5, 7, 8, 9) is backed by one of the sources below, each checked against a primary or "
    "authoritative source (the originating agency, standards body, publisher, or peer-reviewed venue) "
    "while preparing this report.")

add_heading(doc, "Methods", level=2)
add_numbered(doc, [
    "A. N. Angelopoulos and S. Bates. \"Conformal Prediction: A Gentle Introduction.\" Foundations "
    "and Trends® in Machine Learning, 16(4):494–591, 2023. DOI: 10.1561/2200000101. "
    "(Preprint: arXiv:2107.07511, 2021.)",
    "Y. Romano, M. Sesia, and E. Candès. \"Classification with Valid and Adaptive Coverage.\" "
    "Advances in Neural Information Processing Systems 33 (NeurIPS 2020).",
    "V. Vovk. \"Conditional Validity of Inductive Conformal Predictors.\" Proceedings of the Asian "
    "Conference on Machine Learning (ACML 2012).",
    "H. B. McMahan, E. Moore, D. Ramage, S. Hampson, and B. A. y Arcas. \"Communication-Efficient "
    "Learning of Deep Networks from Decentralized Data.\" Proceedings of AISTATS 2017.",
    "A. Hard et al. \"Federated Learning for Mobile Keyboard Prediction.\" 2018 (arXiv:1811.03604).",
    "J. O. du Terrail et al. \"FLamby: Datasets and Benchmarks for Cross-Silo Federated Learning in "
    "Realistic Healthcare Settings.\" 2022 (arXiv:2210.04620).",
])

add_heading(doc, "Dataset", level=2)
add_numbered(doc, [
    "R. Detrano, A. Jánosi, W. Steinbrunn, M. Pfisterer, J. J. Schmid, S. Sandhu, K. H. Guppy, "
    "S. Lee, and V. Froelicher. \"International Application of a New Probability Algorithm for the "
    "Diagnosis of Coronary Artery Disease.\" American Journal of Cardiology, 64(5):304–310, 1989. "
    "DOI: 10.1016/0002-9149(89)90524-9.",
    "UCI Machine Learning Repository. \"Heart Disease\" dataset (ID 45). "
    "archive.ics.uci.edu/dataset/45/heart+disease.",
])

add_heading(doc, "Regulatory and Data-Access Policy (Sections 3-4)", level=2)
add_numbered(doc, [
    "U.S. Dept. of Health and Human Services, Office for Human Research Protections. \"Federal Policy "
    "for the Protection of Human Subjects (45 CFR 46).\" hhs.gov/ohrp.",
    "U.S. HHS Office for Civil Rights. \"Guidance Regarding Methods for De-identification of Protected "
    "Health Information in Accordance with the HIPAA Privacy Rule.\" "
    "hhs.gov/hipaa/for-professionals/special-topics/de-identification.",
    "U.S. HHS. \"Standards for Privacy of Individually Identifiable Health Information.\" Federal "
    "Register, 28 December 2000 (final rule); general compliance date 14 April 2003.",
    "PhysioNet. \"MIMIC-IV\" data-access requirements (PhysioNet Credentialed Health Data Use "
    "Agreement 1.5.0 and CITI \"Data or Specimens Only Research\" training). "
    "physionet.org/content/mimiciv.",
    "U.S. FDA. \"Artificial Intelligence in Software as a Medical Device\" and the AI/ML-Based SaMD "
    "Action Plan. fda.gov/medical-devices/software-medical-device-samd.",
    "Office of the National Coordinator for Health Information Technology. \"Cures Act Final Rule.\" "
    "healthit.gov/regulations/cures-act-final-rule; Federal Register, 1 May 2020.",
    "ONC Recognized Coordinating Entity (The Sequoia Project). \"Trusted Exchange Framework and "
    "Common Agreement (TEFCA).\" rce.sequoiaproject.org/tefca.",
    "Regulation (EU) 2025/327 of the European Parliament and of the Council on the European Health "
    "Data Space. Official Journal of the European Union, 5 March 2025.",
])

page_break(doc)

# ==========================================================================
# APPENDIX
# ==========================================================================
add_heading(doc, "Appendix: Full Figure Index", level=1)
add_para(doc,
    "All 33 figures referenced in this report, in file order, with the script or notebook that "
    "generates each one. Every figure is reproducible by running the listed command from the "
    "repository root.")

appendix_rows = [
    ("figures/01_site_overview.png", "scripts/run_end_to_end_pipeline.py", "Fig. 7.1"),
    ("figures/02_missingness.png", "scripts/run_end_to_end_pipeline.py", "Fig. 5.1"),
    ("figures/03_chol_distributions.png", "scripts/run_end_to_end_pipeline.py", "Fig. 5.2"),
    ("figures/04_js_divergence.png", "scripts/run_end_to_end_pipeline.py", "Fig. 5.6"),
    ("figures/05_domain_auc.png", "scripts/run_end_to_end_pipeline.py", "Fig. 5.5"),
    ("figures/06_fed_curves.png", "scripts/run_end_to_end_pipeline.py", "Fig. 9.1"),
    ("figures/07_calibration_scores.png", "scripts/run_end_to_end_pipeline.py", "Fig. 8.4"),
    ("figures/08_transfer_matrix.png", "scripts/run_end_to_end_pipeline.py", "Fig. 7.4"),
    ("figures/09_coverage_by_site.png", "scripts/run_end_to_end_pipeline.py", "Fig. 7.3"),
    ("figures/10_set_sizes.png", "scripts/run_end_to_end_pipeline.py", "Fig. 8.7"),
    ("figures/11_coverage_beta.png", "scripts/run_end_to_end_pipeline.py", "Fig. 8.5"),
    ("figures/eda/e01_target_distribution.png", "scripts/run_exploratory_data_analysis.py", "Fig. 5.3"),
    ("figures/eda/e02_target_by_site.png", "scripts/run_exploratory_data_analysis.py", "Fig. 5.4"),
    ("figures/eda/e03_continuous_grid.png", "scripts/run_exploratory_data_analysis.py", "Fig. 5.7"),
    ("figures/eda/e04_thalach_by_site.png", "scripts/run_exploratory_data_analysis.py", "Fig. 5.8"),
    ("figures/eda/e05_categorical_grid.png", "scripts/run_exploratory_data_analysis.py", "Fig. 5.9"),
    ("figures/eda/e06_correlation.png", "scripts/run_exploratory_data_analysis.py", "Fig. 5.10"),
    ("figures/cp/c01_class_distribution.png", "scripts/run_chest_pain_pipeline.py", "Fig. 5.11"),
    ("figures/cp/c02_fed_curves.png", "scripts/run_chest_pain_pipeline.py", "Fig. 9.2"),
    ("figures/cp/c03_calibration_scores.png", "scripts/run_chest_pain_pipeline.py", "Fig. 9.7"),
    ("figures/cp/c04_class_conditional.png", "scripts/run_chest_pain_pipeline.py", "Fig. 8.8"),
    ("figures/cp/c05_transfer_matrix.png", "scripts/run_chest_pain_pipeline.py", "Fig. 9.4"),
    ("figures/cp/c06_coverage_by_site.png", "scripts/run_chest_pain_pipeline.py", "Fig. 9.5"),
    ("figures/cp/c07_set_sizes.png", "scripts/run_chest_pain_pipeline.py", "Fig. 9.6"),
    ("figures/paper/fig01_prediction_set_examples.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.1"),
    ("figures/paper/fig02_conformal_illustration.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.2"),
    ("figures/paper/fig04_adaptive_prediction_sets.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.3"),
    ("figures/paper/fig06_conformalized_quantile_regression.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.9"),
    ("figures/paper/fig08_uncertainty_scalar.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.10"),
    ("figures/paper/fig09_conformalized_bayes.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.11"),
    ("figures/paper/fig10_coverage_notions.png", "scripts/run_paper_figure_recreations.py", "Fig. 7.2"),
    ("figures/paper/fig11_coverage_distribution.png", "scripts/run_paper_figure_recreations.py", "Fig. 8.6"),
    ("figures/compare/task_comparison.png", "scripts/compare_prediction_tasks.py", "Fig. 9.3"),
]
add_table(doc, ["File", "Generated by", "Referenced as"], appendix_rows)

add_para(doc,
    "Reproduce everything: python scripts/run_end_to_end_pipeline.py && "
    "python scripts/run_chest_pain_pipeline.py && "
    "python scripts/run_exploratory_data_analysis.py && "
    "python scripts/run_paper_figure_recreations.py && "
    "python scripts/compare_prediction_tasks.py",
    italic=True, size=9.5, color=MUTED)

# ---- Save --------------------------------------------------------------
doc.save(OUT)
print(f"\nReport written to {os.path.abspath(OUT)}")
